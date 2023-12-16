import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from scraping_courses import cleaned_course_names,save_document
product_urls = [
"https://dentistup.tn/produit/interception-en-omni-pratique-2",
"https://dentistup.tn/produit/implantologie-immediate",
"https://dentistup.tn/produit/coming-soon-la-piezographie-mandibulaire",
"https://dentistup.tn/produit/prothese-sur-implants",
"https://dentistup.tn/produit/new-la-chirurgie-de-pose-dimplant",
]
i=0
for product_url in product_urls:
    product_response = requests.get(product_url)
    if product_response.status_code == 200:
        product_soup = BeautifulSoup(product_response.content, "html.parser")
        div_element = product_soup.find("div", class_="woocommerce-product-details__short-description")
        price_element = product_soup.find("span", class_="woocommerce-Price-amount")
        categories_element = product_soup.find("span", class_="posted_in")

        if div_element and price_element and categories_element:
            description = div_element.get_text().strip()
            price = price_element.get_text()
            categories = [a.get_text() for a in categories_element.find_all("a")]
            categories_text = ", ".join(categories)
            doc = Document()
            doc.add_heading("Product Information", level=1)
            doc.add_paragraph(f"Product URL: {product_url}")
            doc.add_paragraph(f"Description: {description}")
            doc.add_paragraph(f"Price: {price}")
            doc.add_paragraph(f"Categories: {categories_text}")
            filename = f"{product_url.split('/')[-1]}.docx"
            folder_path = f"courses//{cleaned_course_names[i]}"
            save_document(doc, folder_path, filename)
            print(f"Saved information for {product_url} in {filename}")
            i+=1
        else:
            print("Product information not found on", product_url)
    else:
        print("Failed to fetch", product_url)

############### only one website
product_urls = ["https://dentistup.tn/produit/pack-implantologie-dentaire/"]

for product_url in product_urls:
    product_response = requests.get(product_url)
    if product_response.status_code == 200:
        product_soup = BeautifulSoup(product_response.content, "html.parser")
        price_element = product_soup.find("p", class_="price")
        price_amount = price_element.find("span", class_="woocommerce-Price-amount").get_text()
        price_currency = price_element.find("span", class_="woocommerce-Price-currencySymbol").get_text()
        price = f"{price_amount} {price_currency}"
        categories_element = product_soup.find("span", class_="posted_in")
        categories = [a.get_text() for a in categories_element.find_all("a")]
        categories_text = ", ".join(categories)
        doc = Document()
        doc.add_heading("Product Information", level=1)
        doc.add_paragraph(f"Product URL: {product_url}")
        doc.add_paragraph(f"Price: {price}")
        doc.add_paragraph(f"Categories: {categories_text}")
        filename = f"{product_url.split('/')[-2]}.docx"
        folder_path = "courses//pack implantologie dentaire"
        save_document(doc, folder_path, filename)
        print(f"Saved information for {product_url} in {filename}")
